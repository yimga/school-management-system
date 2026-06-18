#!/usr/bin/env python3
"""Generate RunMyCampus self-host infrastructure guide as a Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "RunMyCampus_Self_Host_Infrastructure_Guide.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build() -> Document:
    doc = Document()

    title = doc.add_heading("RunMyCampus — Self-Host Infrastructure Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Hardware, software, and service requirements for hosting on your own servers\n"
        "Generated from the RunMyCampus codebase — June 2026"
    )
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph()

    add_heading(doc, "1. Architecture at a Glance", 1)
    add_para(
        doc,
        "RunMyCampus is not a single app process. Production topology (mirrored in "
        "deploy/selfhost/docker-compose.yml and render.yaml) includes:",
    )
    add_bullets(
        doc,
        [
            "Reverse proxy (Caddy / Nginx / Traefik) — TLS termination",
            "Django web (Gunicorn) — HTTP, portals, admin, API",
            "Celery worker — async tasks (email, provisioning, webhooks)",
            "Celery beat — scheduled jobs",
            "PostgreSQL 16 — primary database (schema-per-tenant)",
            "Valkey / Redis — cache, sessions, Celery broker",
            "Ollama (optional) — local AI inference",
            "MinIO or S3-compatible storage (optional) — durable media",
        ],
    )

    add_heading(doc, "2. Required Services by Role", 1)

    add_heading(doc, "2.1 Web Service (Required)", 2)
    add_bullets(
        doc,
        [
            "Software: Python 3.12, Node 20 (build-time), Gunicorn",
            "OS deps: libpq, build tools; WeasyPrint needs libpango, libcairo, libgdk-pixbuf for PDFs",
            "Ports: App listens on 10000 in self-host compose; public 443/80 on reverse proxy",
            "Env: SECRET_KEY, ALLOWED_HOSTS, CSRF_TRUSTED_ORIGINS, MULTI_TENANT_BASE_DOMAIN, USE_DJANGO_TENANTS=1",
        ],
    )
    add_table(
        doc,
        ["Scale", "RAM", "CPU", "Notes"],
        [
            ["Dev / 1 school pilot", "4 GB", "2 cores", "SQLite OK locally; not for multi-tenant prod"],
            ["Small prod (1–5 schools)", "8 GB", "4 cores", "2 Gunicorn workers × 4 threads, Celery concurrency 2"],
            ["Standard prod", "16 GB", "4–8 cores", "Matches Render Standard tier"],
            ["Multi-school / heavy", "32 GB+", "8+ cores", "high_capacity_edge profile in codebase"],
        ],
    )

    add_heading(doc, "2.2 Database — PostgreSQL (Required for Production)", 2)
    add_bullets(
        doc,
        [
            "Software: PostgreSQL 16 + optional pgvector extension (AI memory / RAG)",
            "Storage: SSD strongly recommended; plan 50 GB+ for growth",
            "Backups: Daily pg_dump of the pgdata volume — critical for home hosting",
            "Dedicated DB box minimum: 4 GB RAM, 2 cores, 100 GB SSD",
        ],
    )

    add_heading(doc, "2.3 Cache + Queue Broker — Valkey/Redis (Strongly Recommended)", 2)
    add_bullets(
        doc,
        [
            "REDIS_URL → cache + sessions (moves sessions off Postgres)",
            "CELERY_BROKER_URL → Celery task queue (only if worker is running)",
            "RAM: 512 MB–2 GB; small persistent volume for AOF",
            "Without it: app still runs; tasks run inline; sessions hit Postgres",
        ],
    )

    add_heading(doc, "2.4 Celery Worker (Required for Async Operations)", 2)
    add_bullets(
        doc,
        [
            "Handles: signup provisioning, welcome email, webhooks, heavy reports",
            "RAM: 1–2 GB on top of web; CPU: 2 cores shared OK at small scale",
            "Worker must have same EMAIL_* environment variables as web",
        ],
    )

    add_heading(doc, "2.5 Celery Beat (Required for Schedules)", 2)
    add_bullets(
        doc,
        [
            "Handles: weekly audit checks, key rotation watches, meal-balance sweeps, etc.",
            "RAM: 256–512 MB; can run on same host as worker",
        ],
    )

    add_heading(doc, "2.6 Reverse Proxy + TLS (Required for Real Use)", 2)
    add_bullets(
        doc,
        [
            "Not included in docker-compose — add Caddy, Nginx, or Traefik in front of web:10000",
            "Must terminate HTTPS for apex (yourdomain.com) and wildcard (*.yourdomain.com)",
            "See docs/DEPLOYMENT_SSL_CDN.md and docs/SELF_HOST_MIGRATION.md",
        ],
    )

    add_heading(doc, "2.7 AI Inference — Ollama (Optional)", 2)
    add_para(doc, "Set RMC_DEPLOYMENT_PROFILE=edge for self-hosted AI.")
    add_table(
        doc,
        ["Component", "Where", "Specs"],
        [
            ["Ollama", "Same LAN server or dedicated AI box", "Separate from Django if using large models"],
            ["Small model (qwen2.5:1.5b)", "Pilot", "8 GB RAM min, 16 GB recommended, 8 GB free disk"],
            ["Larger model (llama3.1:8b)", "Production-quality local", "16 GB min, 32 GB recommended, 12 GB free disk"],
            ["Without Ollama", "—", "AI_ALLOW_RULES_FALLBACK=1 — KB/rules answers, no live LLM"],
        ],
    )
    add_para(
        doc,
        "Important: AI runs on the server, not teacher laptops. "
        "OLLAMA_ENDPOINT=http://127.0.0.1:11434 or use docker-compose.ollama.yml sidecar.",
    )

    add_heading(doc, "2.8 Object / Media Storage (Required for Durability)", 2)
    add_bullets(
        doc,
        [
            "Default: local filesystem (MEDIA_ROOT=media/) — back up media/ regularly",
            "Recommended prod: MinIO (self-host, S3-compatible) on same or separate box",
            "Alternative: Cloudflare R2 (10 GB free tier)",
            "Configure: MEDIA_STORAGE_BACKEND=s3 + AWS_S3_ENDPOINT_URL",
        ],
    )

    add_heading(doc, "2.9 Email (Required for Signup, Alerts, Password Reset)", 2)
    add_table(
        doc,
        ["Option", "Home server notes"],
        [
            ["Brevo SMTP (free ~300/day)", "Easiest — no port 25 required"],
            ["Postal (self-host)", "Needs host allowing outbound port 25 + SPF/DKIM/DMARC"],
        ],
    )
    add_para(doc, "Both web and worker services need EMAIL_HOST, EMAIL_HOST_USER, EMAIL_HOST_PASSWORD.")

    add_heading(doc, "2.10 Optional Services (Feature-Dependent)", 2)
    add_table(
        doc,
        ["Service", "Purpose", "Self-host option"],
        [
            ["Collabora", "In-browser document editing", "docker-compose.collabora.yml"],
            ["Meilisearch / Typesense", "Full-text search", "docs/SOVEREIGN_STACK.md"],
            ["OpenSearch", "Log/search at scale", "When OPENSEARCH_DSN set"],
            ["Daphne / ASGI", "WebSockets / SSE streaming", "docs/SSE_DAPHNE_DEPLOYMENT.md"],
            ["Prometheus", "Metrics at /metrics/", "OBSERVABILITY_METRICS_BACKEND=prometheus-client"],
            ["GlitchTip / Sentry", "Error tracking", "Optional SENTRY_DSN"],
            ["HashiCorp Vault", "Secrets management", "Optional; Fernet env keys work without it"],
        ],
    )

    add_heading(doc, "3. Recommended Home-Server Layouts", 1)

    add_heading(doc, "3.1 Option A — Single All-in-One Box (Simplest)", 2)
    add_para(doc, "One machine runs: Web + Worker + Beat + PostgreSQL + Valkey + Ollama (optional) + MinIO (optional) + Reverse proxy.")
    add_table(
        doc,
        ["Schools / users", "RAM", "CPU", "Storage"],
        [
            ["Lab / 1 school, no AI", "8 GB", "4 cores", "256 GB SSD"],
            ["1–5 schools + small AI", "16 GB", "6–8 cores", "512 GB SSD"],
            ["5+ schools + local AI", "32 GB+", "8+ cores", "1 TB NVMe"],
        ],
    )
    add_para(doc, "Quick start commands:", bold=True)
    add_bullets(
        doc,
        [
            "cp deploy/selfhost/.env.example deploy/selfhost/.env",
            "Edit secrets and domain settings in .env",
            "docker compose -f deploy/selfhost/docker-compose.yml up -d --build",
            "Optional AI: docker compose -f docker-compose.ollama.yml up -d",
        ],
    )
    add_para(doc, "Reference docs: docs/SELF_HOST_MIGRATION.md, deploy/selfhost/docker-compose.yml")

    add_heading(doc, "3.2 Option B — Split Across 2–3 Home Servers (Healthier)", 2)
    add_table(
        doc,
        ["Server", "Role", "Specs"],
        [
            ["Server 1 — App", "Web, Celery worker, beat, reverse proxy, MinIO (optional)", "8–16 GB RAM, 4–8 cores, 256 GB SSD"],
            ["Server 2 — Data", "PostgreSQL 16 (+ pgvector), nightly backups", "8–16 GB RAM, 4 cores, 500 GB–1 TB SSD/NVMe"],
            ["Server 3 — AI (optional)", "Ollama only", "16–32 GB RAM, GPU optional, 256 GB SSD"],
        ],
    )
    add_para(doc, "Point DATABASE_URL and REDIS_URL at Server 2 / Valkey on DB or app server.")

    add_heading(doc, "4. External Dependencies (Cannot Fully Self-Host)", 1)
    add_table(
        doc,
        ["Need", "Why", "Typical approach"],
        [
            ["SMS", "Carrier termination", "AfricasTalking / Twilio (pay per message)"],
            ["Card payments", "PCI / payment gateways", "Stripe / Paystack (per-transaction fees)"],
            ["Domain + DNS", "Registrar requirement", "~$10–15/year"],
            ["Outbound email at scale", "Port 25 blocked on most home ISPs", "Brevo SMTP or relay VPS (~€4/mo)"],
        ],
    )
    add_para(doc, "AI, DB, cache, queue, and core app can stay 100% on your hardware.")

    add_heading(doc, "5. Minimum vs Recommended Summary", 1)
    add_table(
        doc,
        ["Tier", "Servers", "RAM total", "What you get"],
        [
            ["Minimum viable", "1", "8 GB", "Web + Postgres + Valkey + worker + beat; rules-only AI; Brevo email"],
            ["Recommended home prod", "1–2", "16–32 GB", "Full async stack, local Ollama (small model), MinIO media, backups"],
            ["Comfortable multi-school", "2–3", "32–64 GB", "Split DB + dedicated AI box, larger models, Collabora optional"],
        ],
    )

    add_heading(doc, "6. First-Time Setup Checklist", 1)
    steps = [
        "Linux host (Ubuntu 22.04/24.04 or Debian) with Docker + Compose",
        "Copy deploy/selfhost/.env.example → .env; set SECRET_KEY, POSTGRES_PASSWORD, domains",
        "docker compose -f deploy/selfhost/docker-compose.yml up -d --build",
        "Reverse proxy + Let's Encrypt for apex + *.yourdomain.com",
        "Migrations run automatically on first boot via entrypoint.web.sh",
        "Seed admin: docker compose ... exec web python manage.py ensure_superuser",
        "Configure SMTP on both web and worker services",
        "Optional: Ollama + RMC_DEPLOYMENT_PROFILE=edge",
        "Schedule Postgres + media backups",
        "Verify: GET https://yourdomain.com/health/ → database, redis, celery broker configured",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    add_heading(doc, "7. Deployment Profiles (AI)", 1)
    add_table(
        doc,
        ["Profile", "Typical host", "AI behavior"],
        [
            ["online (default)", "Cloud SaaS", "LiteLLM / cloud API → Ollama → rules fallback"],
            ["edge", "School LAN hub", "Ollama → rules fallback"],
            ["hybrid", "Cloud + optional hub", "Cloud AI on Render; hub Ollama when served from hub"],
        ],
    )
    add_para(doc, "See docs/AI_DEPLOYMENT_POSTURE.md and docs/LOCAL_HUB_MODE.md for full details.")

    add_heading(doc, "8. Cost-Aware Adoption Stages", 1)
    add_para(doc, "From docs/PRODUCTION_INFRA_FOUNDATION.md — adopt in stages:")
    add_bullets(
        doc,
        [
            "Tier 1: Upgrade Postgres first (biggest single win for performance)",
            "Tier 2: Add Valkey for cache + sessions (REDIS_URL only — no CELERY_BROKER_URL yet)",
            "Tier 3: Add Celery worker + set CELERY_BROKER_URL (complete async fix)",
            "Rule: Only set CELERY_BROKER_URL when a worker is actually running",
        ],
    )

    add_heading(doc, "9. Key Repository References", 1)
    add_bullets(
        doc,
        [
            "deploy/selfhost/docker-compose.yml — full self-host stack",
            "deploy/selfhost/.env.example — environment template",
            "docs/SELF_HOST_MIGRATION.md — Render → self-host migration",
            "docs/SELF_HOSTING_AND_EXTERNAL_DEPENDENCIES_BACKLOG.md — FOSS vs paid deps",
            "docs/PRODUCTION_INFRA_FOUNDATION.md — Postgres + Valkey + Celery rationale",
            "docs/DEPLOYMENT_FULL.md — production deploy reference",
            "docs/OLLAMA_OPERATIONS_AND_UPDATES.md — local AI setup",
            "docs/SOVEREIGN_STACK.md — open-source technology choices",
            "render.yaml — production topology reference (Render SaaS)",
        ],
    )

    add_heading(doc, "10. Important Caveat", 1)
    add_para(
        doc,
        "docs/SELF_HOST_MIGRATION.md marks the self-host compose stack as scaffold / "
        "not yet proven against full production data. Safe for home lab and staging; "
        "test with a copy of production data before cutover.",
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("RunMyCampus — school-management-system repository")
    fr.font.size = Pt(9)
    fr.italic = True

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
